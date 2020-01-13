from PyQt5 import QtCore, QtGui, QtWidgets
from Module.sql import *
from Module.pelanggan_new import *
import os,datetime
from docx import Document
from docx.shared import Mm
from docx.shared import Pt
from docx.shared import Length

class Transaksi():
    def __init__(self, w):
        self.pelanggan = w.mysql
        self.parent = w
        self.transaksiWidget = QtWidgets.QWidget(w)
        self.transaksiWidget.setEnabled(True)
        self.transaksiWidget.setGeometry(QtCore.QRect(230, 0, 1141, 771))
        self.transaksiWidget.setStyleSheet("")
        self.transaksiWidget.setObjectName("transaksiWidget")
        self.label_gender = QtWidgets.QLabel(self.transaksiWidget)
        self.label_gender.setGeometry(QtCore.QRect(440, 112, 61, 25))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_gender.setFont(font)
        self.label_gender.setObjectName("label_gender")
        self.id = QtWidgets.QLineEdit(self.transaksiWidget)
        self.id.setGeometry(QtCore.QRect(518, 52, 371, 25))
        font = QtGui.QFont()
        font.setPointSize(12)
        # self.id.setFont(font)
        self.id.setPlaceholderText("Masukkan idPelanggan atau scan kartu")
        self.id.setClearButtonEnabled(True)
        self.id.setObjectName("id")
        self.id.returnPressed.connect(self.search)
        # self.clear = qt
        self.alamat = QtWidgets.QLineEdit(self.transaksiWidget)
        self.alamat.setGeometry(QtCore.QRect(518, 142, 371, 25))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.alamat.setFont(font)
        self.alamat.setClearButtonEnabled(True)
        self.alamat.setObjectName("alamat")
        # self.alamat.setReadOnly(True)
        self.horizontal_line = QtWidgets.QFrame(self.transaksiWidget)
        self.horizontal_line.setGeometry(QtCore.QRect(40, 290, 1061, 20))
        self.horizontal_line.setStyleSheet("")
        self.horizontal_line.setFrameShape(QtWidgets.QFrame.HLine)
        self.horizontal_line.setFrameShadow(QtWidgets.QFrame.Sunken)
        self.horizontal_line.setObjectName("horizontal_line")
        self.nomor_hp = QtWidgets.QLineEdit(self.transaksiWidget)
        self.nomor_hp.setGeometry(QtCore.QRect(518, 172, 371, 25))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.nomor_hp.setFont(font)
        # self.nomor_hp.setClearButtonEnabled(True)
        self.nomor_hp.setObjectName("nomor_hp")
        self.nomor_hp.setReadOnly(True)
        self.nama = QtWidgets.QLineEdit(self.transaksiWidget)
        self.nama.setGeometry(QtCore.QRect(518, 82, 371, 25))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.nama.setFont(font)
        self.nama.setClearButtonEnabled(True)
        self.nama.setObjectName("nama")
        # self.nama.setReadOnly(True)
        self.label_alamat = QtWidgets.QLabel(self.transaksiWidget)
        self.label_alamat.setGeometry(QtCore.QRect(440, 142, 51, 25))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_alamat.setFont(font)
        self.label_alamat.setObjectName("label_alamat")
        self.gender = QtWidgets.QLineEdit(self.transaksiWidget)
        self.gender.setGeometry(QtCore.QRect(518, 112, 371, 25))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.gender.setFont(font)
        self.gender.setReadOnly(True)
        self.img = QtWidgets.QLabel(self.transaksiWidget)
        pixmap = QtGui.QPixmap("placeholder.png")
        self.img.setPixmap(pixmap.scaled(200, 170, QtCore.Qt.KeepAspectRatio))
        self.img.setObjectName("img")
        self.img.setGeometry(QtCore.QRect(180, 20, 211, 241))
        self.label_nama = QtWidgets.QLabel(self.transaksiWidget)
        self.label_nama.setGeometry(QtCore.QRect(440, 82, 41, 20))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_nama.setFont(font)
        self.label_nama.setObjectName("label_nama")
        self.label_nomorHP = QtWidgets.QLabel(self.transaksiWidget)
        self.label_nomorHP.setGeometry(QtCore.QRect(440, 172, 51, 25))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_nomorHP.setFont(font)
        self.label_nomorHP.setObjectName("label_nomorHP")
        self.label_id = QtWidgets.QLabel(self.transaksiWidget)
        self.label_id.setGeometry(QtCore.QRect(440, 52, 31, 25))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_id.setFont(font)
        self.label_id.setObjectName("label_id")
        self.scrollArea = QtWidgets.QScrollArea(self.transaksiWidget)
        self.scrollArea.setGeometry(QtCore.QRect(20, 320, 1101, 311))
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.tabelView = QtWidgets.QWidget()
        self.tabelView.setGeometry(QtCore.QRect(0, 0, 1099, 309))
        self.tabelView.setObjectName("tabelView")
        self.tabelTransaksi = QtWidgets.QTableWidget(self.tabelView)
        self.tabelTransaksi.setGeometry(QtCore.QRect(0, 0, 1101, 311))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Expanding)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.tabelTransaksi.sizePolicy().hasHeightForWidth())
        self.tabelTransaksi.setSizePolicy(sizePolicy)
        self.tabelTransaksi.setFocusPolicy(QtCore.Qt.ClickFocus)
        self.tabelTransaksi.setSizeAdjustPolicy(QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.tabelTransaksi.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self.tabelTransaksi.setSelectionMode(QtWidgets.QAbstractItemView.ExtendedSelection)
        self.tabelTransaksi.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.tabelTransaksi.setObjectName("tabelTransaksi")
        self.tabelTransaksi.setColumnCount(6)
        self.tabelTransaksi.setRowCount(0)
        item = QtWidgets.QTableWidgetItem()
        self.tabelTransaksi.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.tabelTransaksi.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.tabelTransaksi.setHorizontalHeaderItem(2, item)
        item = QtWidgets.QTableWidgetItem()
        self.tabelTransaksi.setHorizontalHeaderItem(3, item)
        item = QtWidgets.QTableWidgetItem()
        self.tabelTransaksi.setHorizontalHeaderItem(4, item)
        item = QtWidgets.QTableWidgetItem()
        self.tabelTransaksi.setHorizontalHeaderItem(5, item)
        self.tabelTransaksi.horizontalHeader().setCascadingSectionResizes(True)
        header = self.tabelTransaksi.horizontalHeader()
        header.setSectionResizeMode(1, QtWidgets.QHeaderView.Stretch)

        self.scrollArea.setWidget(self.tabelView)
        self.saldo2 = QtWidgets.QLineEdit(self.transaksiWidget)
        self.saldo2.setGeometry(QtCore.QRect(790, 240, 281, 41))
        self.saldo2.setStyleSheet("background-color: #00c853;\n"
            "border: 2px solid #00c853;\n"
            "color:white;\n"
            "padding: 0 10px;\n"
            "font-size: 18pt;")
        self.cancel = QtWidgets.QPushButton(self.transaksiWidget)
        self.cancel.setText("Cancel")
        self.cancel.setFont(font)
        self.cancel.setGeometry(100, 710, 131, 31)
        self.cancel.setStyleSheet("QPushButton {\n"
                "    background-color: red;\n"
                "    border:none;\n"
                "    color:white;\n"
                "    border-radius: 6px;\n"
                "}\n"
                "QPushButton:pressed {\n"
                "    background-color: white;\n"
                "    border: 2px solid red;\n"
                "    color: red;\n"
                "}")
        self.cancel.clicked.connect(self.dont)
        self.saldo2.setAlignment(QtCore.Qt.AlignCenter)
        self.saldo2.setReadOnly(True)
        self.saldo2.setObjectName("saldo")
        self.label_saldo = QtWidgets.QLabel(self.transaksiWidget)
        self.label_saldo.setGeometry(QtCore.QRect(906, 207, 51, 20))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.label_saldo.setFont(font)
        self.label_saldo.setStyleSheet("color: #00c853;")
        self.label_saldo.setObjectName("label_saldo")
        self.point = QtWidgets.QLineEdit(self.transaksiWidget)
        self.point.setEnabled(True)
        self.point.setGeometry(QtCore.QRect(580, 240, 171, 41))
        self.point.setStyleSheet("background-color: #ff8f00;\n"
            "border: 2px solid #ff8f00;\n"
            "color:white;\n"
            "padding: 0 10px;\n"
            "font-size: 18pt;")
        self.point.setAlignment(QtCore.Qt.AlignCenter)
        self.point.setReadOnly(True)
        self.point.setObjectName("point")
        self.label_point = QtWidgets.QLabel(self.transaksiWidget)
        self.label_point.setGeometry(QtCore.QRect(642, 207, 51, 20))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.label_point.setFont(font)
        self.label_point.setStyleSheet("color: #ff8f00;")
        self.label_point.setObjectName("label_point")
        self.total = QtWidgets.QLineEdit(self.transaksiWidget)
        self.total.setGeometry(QtCore.QRect(900, 650, 221, 41))
        self.total.setStyleSheet("font-size: 14pt;\n"
            "padding: 0 10px;")
        self.total.setAlignment(QtCore.Qt.AlignCenter)
        self.total.setReadOnly(True)
        self.total.setObjectName("total")
        self.label_total = QtWidgets.QLabel(self.transaksiWidget)
        self.label_total.setGeometry(QtCore.QRect(840, 660, 51, 20))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.label_total.setFont(font)
        self.label_total.setStyleSheet("")
        self.label_total.setObjectName("label_total")
        self.prosesBtn = QtWidgets.QPushButton(self.transaksiWidget)
        self.prosesBtn.setGeometry(QtCore.QRect(980, 710, 131, 31))
        self.prosesBtn.setStyleSheet("QPushButton {\n"
            "    background: #2962ff;\n"
            "    border-radius: 6px;\n"
            "    border: none;\n"
            "    color: white;\n"
            "    outline: none;\n"
            "}\n"
            "\n"
            "QPushButton:pressed {\n"
            "    border: 2px solid #2962ff;\n"
            "    background: white;\n"
            "    color: #2962ff;\n"
            "}")
        self.prosesBtn.setObjectName("prosesBtn")
        self.prosesBtn.clicked.connect(self.go)
        self.hapusBtn = QtWidgets.QPushButton(self.transaksiWidget)
        self.hapusBtn.setGeometry(QtCore.QRect(830, 710, 131, 31))
        self.hapusBtn.setStyleSheet("QPushButton {\n"
            "    background: #ff3d00;\n"
            "    border-radius: 6px;\n"
            "    border: none;\n"
            "    color: white;\n"
            "    outline: none;\n"
            "}\n"
            "\n"
            "QPushButton:pressed {\n"
            "    border: 2px solid #ff3d00;\n"
            "    background: white;\n"
            "    color: #ff3d00;\n"
            "}")
        self.hapusBtn.setFlat(False)
        self.hapusBtn.setObjectName("hapusBtn")
        self.hapusBtn.clicked.connect(self.deleteProduct)
        self.comboBox = QtWidgets.QComboBox(self.transaksiWidget)
        self.comboBox.setGeometry(QtCore.QRect(20, 650, 211, 41))
        self.comboBox.setStyleSheet("font-size: 14pt;\n"
            "padding: 0 10px")
        self.comboBox.setObjectName("comboBox")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")

        self.retranslateUi(w)
        QtCore.QMetaObject.connectSlotsByName(self.transaksiWidget)
        self.transaksiWidget.close()

    def retranslateUi(self, w):
        _translate = QtCore.QCoreApplication.translate
        self.label_gender.setText(_translate("transaksiWidget", "Gender"))
        self.label_alamat.setText(_translate("transaksiWidget", "Alamat"))
        self.label_nama.setText(_translate("transaksiWidget", "Nama"))
        self.label_nomorHP.setText(_translate("transaksiWidget", "No. HP"))
        self.label_id.setText(_translate("transaksiWidget", "ID"))
        item = self.tabelTransaksi.horizontalHeaderItem(0)
        item.setText(_translate("transaksiWidget", "ID Produk"))
        item = self.tabelTransaksi.horizontalHeaderItem(1)
        item.setText(_translate("transaksiWidget", "Nama Produk"))
        item = self.tabelTransaksi.horizontalHeaderItem(2)
        item.setText(_translate("transaksiWidget", "Harga"))
        item = self.tabelTransaksi.horizontalHeaderItem(3)
        item.setText(_translate("transaksiWidget", "Qty"))
        item = self.tabelTransaksi.horizontalHeaderItem(4)
        item.setText(_translate("transaksiWidget", "Diskon%"))
        item = self.tabelTransaksi.horizontalHeaderItem(5)
        item.setText(_translate("transaksiWidget", "Jumlah"))
        self.saldo2.setText(_translate("transaksiWidget", "Rp 0"))
        self.label_saldo.setText(_translate("transaksiWidget", "Saldo"))
        self.point.setText(_translate("transaksiWidget", "0"))
        self.label_point.setText(_translate("transaksiWidget", "Point"))
        self.total.setText(_translate("transaksiWidget", "Rp 0"))
        self.label_total.setText(_translate("transaksiWidget", "Total"))
        self.prosesBtn.setText(_translate("transaksiWidget", "Proses"))
        self.hapusBtn.setText(_translate("transaksiWidget", "Hapus"))
        self.comboBox.setItemText(0, _translate("transaksiWidget", "Line 1"))
        self.comboBox.setItemText(1, _translate("transaksiWidget", "Line 2"))
        self.comboBox.setItemText(2, _translate("transaksiWidget", "Line 3"))

    def show(self):
            self.transaksiWidget.show()
            
    def close(self):
            self.transaksiWidget.close()

    def showDialog(self, msgType: str, title: str, message: str):
        msgBox = QtWidgets.QMessageBox()
        if msgType is "critical":
            msgBox.critical(self, title, message)
        elif msgType is "warning":
            msgBox.warning(self, title, message)
        elif msgType is "information":
            msgBox.information(self, title, message)
        elif msgType is "question":
            return msgBox.question(self, title, message)

    def deleteProduct(self):
        i = 0
        tt = self.tabelTransaksi.currentRow()
        if tt <= -1:
            self.parent.showDialog("information", "Informasi","Tolong pilih item yang hendak dihapus")
            return 
        else:
            self.tabelTransaksi.removeRow(tt) 
        self.count()

    def search(self):
        if self.id.text() == "":
            self.parent.showDialog("warning","warning","tidak ada data yang cocok")
            return
        listPelanggan = self.pelanggan.find(
            "idPelanggan,nama,gender,alamat,kontak,saldo,point", "pelanggan",
            "idPelanggan", self.id.text(), True,
            "idPelanggan ASC")
        if len(listPelanggan) < 1:
            listPelanggan = self.pelanggan.find(
                "idPelanggan,nama,gender,alamat,kontak,saldo,point", "pelanggan",
                "rfid", self.id.text(), True,
                "idPelanggan ASC")
        if len(listPelanggan) < 1:
            self.parent.showDialog("warning","warning","tidak ada data yang cocok")
            return
        if len(listPelanggan) > 1:
            col = 0
            for data in listPelanggan:
                for val in data:
                    if col is 0:
                        val = self.id.setText(str(val))
                        path = QtCore.QDir.currentPath()
                        path = path + '/Module/static'
                        imagepath = path + '/' + self.id.text()+str(".jpg")
                        pixmap = QtGui.QPixmap(imagepath)
                        self.img.setPixmap(pixmap.scaled(200,170, QtCore.Qt.KeepAspectRatio))
                    if col is 1:
                        val = self.nama.setText(str(val))
                    if col is 2:
                        val = self.gender.setText(str(val))
                    if col is 3:
                        val = self.alamat.setText(str(val))
                    if col is 4:
                        val = self.nomor_hp.setText(str(val))
                    if col is 5:
                        val = self.saldo2.setText(str(val))
                    if col is 6:
                        val = self.point.setText(str(val))
                    col = col + 1

    def dont(self):
        self.id.clear()
        self.nama.clear()
        self.gender.clear()
        self.alamat.clear()
        self.nomor_hp.clear()
        self.saldo2.clear()
        self.point.clear()
        self.total.setText("Rp0")
        pixmap = QtGui.QPixmap("placeholder.png")
        self.img.setPixmap(pixmap.scaled(200, 170, QtCore.Qt.KeepAspectRatio))
        row = self.tabelTransaksi.rowCount()
        for i in range (row):
            self.tabelTransaksi.removeRow(0)
            self.tabelTransaksi.removeRow(i)

    def undo(self):
        self.id.clear()
        self.nama.clear()
        self.gender.clear()
        self.alamat.clear()
        self.nomor_hp.clear()
        self.saldo2.clear()
        self.point.clear()
        pixmap = QtGui.QPixmap("placeholder.png")
        self.img.setPixmap(pixmap.scaled(200, 170, QtCore.Qt.KeepAspectRatio))

    def count(self):
        self.ttl =[]
        self.ttlname=[]
        self.ttlz =[]
        self.ttlid =[]
        self.ttlprice=[]
        self.rowz = self.tabelTransaksi.rowCount()
        for i in range (self.rowz):
            itemid = self.tabelTransaksi.item(i,0).text()
            itemname=self.tabelTransaksi.item(i,1).text()
            itemprice=self.tabelTransaksi.item(i,2).text() 
            itemz = self.tabelTransaksi.item(i,3).text()
            item = self.tabelTransaksi.item(i,5).text()
            self.ttl.append(int(item))
            self.ttlprice.append(itemprice)
            self.ttlname.append(itemname)
            self.ttlz.append(int(itemz))
            self.ttlid.append(itemid)
        self.alls = 0
        for x in range (self.rowz):
            self.alls = self.alls + self.ttl[x]
        self.total.setText("Rp " + str(self.alls))
        self.struck()


    def struck(self):
        row = self.tabelTransaksi.rowCount()
        self.records = []
        for i in range (row):
            self.records.append([])
            for j in range (4):
                if j == 0:
                    item = self.tabelTransaksi.item(i,1).text()
                if j == 1:
                    item = self.tabelTransaksi.item(i,3).text()
                if j == 2:
                    item = self.tabelTransaksi.item(i,2).text()
                if j == 3:
                    item = self.tabelTransaksi.item(i,5).text()
                self.records[i].append(item)

    def go(self):
        saldo = self.saldo2.text()
        x = self.total.text()
        price = x[3:]
        row = self.tabelTransaksi.rowCount()
        if row == 0:
            self.parent.showDialog("warning", "warning","anda belum memilih item untuk dibeli")
        if row > 0:
            if self.nama.text() == "":
                self.parent.showDialog("warning", "warning","tidak ada data pelanggan pada form")
                return
            if (int(saldo) < int(price)):
                self.parent.showDialog("warning", "warning","saldo tidak mencukupi silahkan lakukan Top up terlebih dahulu")
                return
            if (int(saldo) >= int(price)):
                self.idt = self.pelanggan.select("id", "transaksi", False, "id DESC")
                if self.idt != None :
                    self.idt = '{:09d}'.format(self.idt[0]+1)
                else:
                    self.idt = '{:09d}'.format(1)
                self.idt = "PAY-" + str(self.idt)
                m = []
                for i in range (self.rowz):
                    keyz = "idProduk = \"{}\""
                    keyz = keyz.format(self.ttlid[i])
                    idp = self.ttlid[i]
                    nam = self.ttlname[i]
                    har = self.ttlprice[i]
                    tot = self.ttl[i]
                    y = self.pelanggan.findCol("stock","daftarbarang",keyz,True)
                    stockambil = self.ttlz[i]
                    for s in range(1):
                        stockawal = y[0][0]
                        stockakhir =  stockawal - stockambil
                        keluar = 0 - stockambil
                        val = "stock =  {}"
                        val = val.format(stockakhir)
                        vals = "\"{}\",\"{}\",{}"
                        vals = vals.format(idp,nam,keluar)
                        self.pelanggan.insertTo("stock","idProduk,nama,jumlah",vals)
                        self.pelanggan.update("daftarbarang",val,keyz)
                        valz = "\"{}\",\"{}\",\"{}\",{},{},\"{}\",\"{}\",\"{}\",{}"
                        valz = valz.format(self.idt,idp,nam,har,stockambil,self.id.text(),self.nama.text(),"beli",tot)
                        self.pelanggan.insertTo("transaksi","idTransaksi,idProduk,produk,harga,jumlah,idPelanggan,pelanggan,jenisTransaksi,total",valz)

                money = int(self.saldo2.text())
                lastmoney = money - self.alls
                val = "saldo= {}"
                val = val.format(lastmoney)
                key = "idPelanggan = \"{}\""
                key = key.format(self.id.text())
                self.pelanggan.update("pelanggan",val,key)
                self.struck()
                self.print()
                self.parent.showDialog("information", "Informasi","Transaksi Berhasil")
                self.dont()
                # QtWidgets.QMessageBox.Yes:
        # print("halo")

    def print(self):
        document = Document()

        #atur paper
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(58)
        section.left_margin = Mm(5.08)
        section.right_margin = Mm(10.16)
        section.top_margin = Mm(5.08)
        section.bottom_margin = Mm(86.868)
        section.space_before = Pt(0)
        section.space_after = Pt(0)
        # document.add_picture('download.png', width=Mm(47.4), height=Mm(24))
        
        #atur font
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(7)
        
        #atur waktu
        t = datetime.datetime.now()
        date = t.strftime("%d/%m/%Y")
        time = t.strftime("%H:%M")
        
        #print date and time and id transaksi
        id_transaksi = self.idt
        p1 = document.add_paragraph()
        p1.alignment = 1
        p1.add_run(date)
        p1.add_run(' ')
        p1.add_run(time)
        p1.add_run(' " ')
        p1.add_run(id_transaksi)
        p1.add_run(' " ')
        
        #print list order id pelanggan
        self.rr = self.id.text()
        id_pelanggan = self.rr
        p2 = document.add_paragraph()
        p2.alignment = 1
        p2.add_run('List Order from')
        p2.add_run(' " ')
        p2.add_run(id_pelanggan)
        p2.add_run(' " ')
        
        #print list permintaan
        much_item = 4
        
        table = document.add_table(rows=1, cols=much_item)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'barang'
        hdr_cells[1].text = 'QTY'
        hdr_cells[2].text = 'harga'
        hdr_cells[3].text = 'total'
        for barang,QTY, harga, total in self.records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(barang)
            row_cells[1].text = str(QTY)
            row_cells[2].text = str(harga)
            row_cells[3].text = str(total)
            
        #garis bawah
        p3 = document.add_paragraph()
        p3.alignment = 3
        p3.add_run('.').underline = True
        p3.add_run('                                                            ').underline = True
        p3.add_run('.').underline = True
        
        #total harga
        self.tt = self.total.text() 
        result = self.tt
        p4 = document.add_paragraph()
        p4.alignment = 3
        p4.add_run('total harga :')
        p4.add_run('               ')
        p4.add_run(result)
        
        path = QtCore.QDir.currentPath()
        path = path + '/Module/Data Struk'
        printpath = path + '/' + "struk_transaksi.docx"

        document.save(printpath)

        os.startfile(printpath,"print")